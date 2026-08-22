"""Facilitator Payment Letter PDF — python-docx fills original DOCX → LibreOffice → PDF.

Template: app/static/templates/docx/payment_letter.docx  (copy of original DOCX)
No Jinja2/docxtpl — runs and table rows are patched directly.
"""

import copy
import io
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.services.documents.signature_block import (
    SignatureParty,
    column_headings,
    decode_signature,
    replace_signature_block,
    template_signature_image,
)

_TEMPLATE = (
    Path(__file__).parent.parent.parent
    / "static" / "templates" / "docx" / "payment_letter.docx"
)

_SOFFICE = (
    r"C:\Program Files\LibreOffice\program\soffice.exe"
    if sys.platform == "win32"
    else "libreoffice"
)


def _libreoffice_to_pdf(docx_bytes: bytes) -> bytes:
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "doc.docx"
        src.write_bytes(docx_bytes)
        subprocess.run(
            [_SOFFICE, "--headless", "--convert-to", "pdf", "--outdir", tmp, str(src)],
            check=True, timeout=60, capture_output=True,
        )
        return (Path(tmp) / "doc.pdf").read_bytes()


def _currency(v) -> str:
    return f"{float(v):,.0f}"


def _replace_aptos(doc: Document) -> None:
    """Replace Aptos font references with Calibri throughout the document.

    Aptos is a Microsoft 365 font unavailable in LibreOffice; without it LibreOffice
    substitutes bold text with a serif fallback. Calibri is visually close and
    ships with Windows in regular, bold, italic, and bold-italic variants.
    """
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for rfont in doc.element.iter(qn("w:rFonts")):
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            key = f"{{{ns}}}{attr}"
            if rfont.get(key) == "Aptos":
                rfont.set(key, "Calibri")


def _wt(run_el, idx: int, text: str) -> None:
    """Set the n-th <w:t> text inside a run element, preserving xml:space."""
    wts = run_el.findall(qn("w:t"))
    if idx < len(wts):
        wts[idx].text = text
        wts[idx].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _ensure_rfonts(run, para) -> None:
    """Copy <w:rFonts> from paragraph pPr.rPr into the run's rPr if missing.

    LibreOffice doesn't apply the OOXML paragraph→run font inheritance, so runs
    that rely on pPr.rPr for their font need an explicit rFonts element.
    """
    run_rpr = run._r.find(qn("w:rPr"))
    if run_rpr is None or run_rpr.find(qn("w:rFonts")) is not None:
        return
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return
    para_rpr = pPr.find(qn("w:rPr"))
    if para_rpr is None:
        return
    rFonts = para_rpr.find(qn("w:rFonts"))
    if rFonts is not None:
        run_rpr.insert(0, copy.deepcopy(rFonts))


def _fill_cell(cell, text: str) -> None:
    """Set text in a table cell, keeping the cell's paragraph/run formatting."""
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = str(text)
        for run in para.runs[1:]:
            run.text = ""
        _ensure_rfonts(para.runs[0], para)
    else:
        run = para.add_run(str(text))
        _ensure_rfonts(run, para)


def _rebuild_signature_block(
    doc: Document,
    paras: list,
    admin_name: str,
    instructor_name: str,
    signed_date: str | None,
    admin_sig_bytes: bytes | None,
    instructor_sig_b64: str | None,
) -> None:
    """Replace P[30]/P[31] — the signature block — with a two-column table.

    The template spreads the block over two paragraphs (P[30] headings + the
    "Name:"/"Date:" rows, P[31] the "Signature:" row with the admin's signature
    anchored to it) and holds the columns apart with counted tab stops measured
    against one pair of names. A longer name wraps the row and pushes the field
    beside it out of column, so the whole block is rebuilt as a table instead —
    see signature_block.py. The admin's signature falls back to the one baked
    into the template when no override is supplied, as it did before.
    """
    template_sig, sig_size = template_signature_image(doc, paras[1])
    left_heading, right_heading = column_headings(
        paras[0].text, ("For SpacePoint FZC", "Facilitator")
    )
    # Per the client's requirement, neither signature date shows until the
    # instructor actually signs — `letter_date` (the letter's own issuance
    # date, a separate concept) is still used in the P[1] header above.
    date = signed_date or ""
    replace_signature_block(
        doc,
        paras,
        SignatureParty(
            heading=left_heading,
            name=admin_name,
            date=date,
            signature=admin_sig_bytes or template_sig,
        ),
        SignatureParty(
            heading=right_heading,
            name=instructor_name,
            date=date,
            signature=decode_signature(instructor_sig_b64),
        ),
        signature_size=sig_size,
    )


def generate_payment_letter_pdf(
    *,
    instructor_name: str,
    reference: str,
    letter_date: str,
    sessions: list[dict],
    addons: list[dict],
    bank: dict | None,
    admin_signatory_name: str,
    admin_signatory_title: str = "",
    admin_signature_bytes: bytes | None = None,
    instructor_signature_b64: str | None = None,
    signed_date: str | None = None,
) -> bytes:
    """Each session dict: {session_date, workshop_description, role, location,
    duration_hours, compensation_aed}. Each addon dict: {description,
    amount_aed, notes}. `bank`: {account_holder_name, bank_name, iban, swift_bic} or None.
    """
    doc = Document(str(_TEMPLATE))
    _replace_aptos(doc)
    paras = doc.paragraphs
    bank = bank or {}

    base_total = sum(s["compensation_aed"] for s in sessions)
    addons_total = sum(a["amount_aed"] for a in addons)
    grand_total = base_total + addons_total

    # P[1]: header block — "Date:", "Facilitator:", "Reference Agreement:"
    # run[3] = ':\n'  (after 'Date')      — first <w:t> is ':' then <w:br/>
    # run[5] = ': \n' (after 'Facilitator') — first <w:t> is ': ' then <w:br/>
    # run[7] = ': Facilitator Agreement '  — plain run, no break
    _wt(paras[1].runs[3]._r, 0, f": {letter_date}")
    _wt(paras[1].runs[5]._r, 0, f": {instructor_name}")
    paras[1].runs[7].text = f": {reference or 'Facilitator Agreement'}"

    # P[9]: Total Optional Add-ons: AED 300  →  run[1] = ': AED 300'
    paras[9].runs[1].text = f": AED {_currency(addons_total)}"

    # P[13]: Total Compensation Payable: AED 2,600  (single bold run)
    paras[13].runs[0].text = f"Total Compensation Payable: AED {_currency(grand_total)}"

    # P[14]: breakdown sentence
    paras[14].runs[0].text = (
        f"Total includes base compensation of AED {_currency(base_total)}"
        f" + optional add-ons of AED {_currency(addons_total)}."
    )

    # P[21-24]: bank details — each has run[0]=bold label, run[1]=value
    paras[21].runs[1].text = f": {bank.get('account_holder_name', '')}"
    paras[22].runs[1].text = f": {bank.get('bank_name', '')}"
    paras[23].runs[1].text = f": {bank.get('iban', '')}"
    paras[24].runs[1].text = f" (if applicable): {bank.get('swift_bic', '')}"

    # --- Table 0: Sessions ---
    tbl0 = doc.tables[0]
    orig_tr = tbl0.rows[1]._tr
    row_template = copy.deepcopy(orig_tr)
    tbl0._tbl.remove(orig_tr)
    for s in sessions:
        new_tr = copy.deepcopy(row_template)
        tbl0._tbl.append(new_tr)
        row = tbl0.rows[-1]
        _fill_cell(row.cells[0], s["session_date"])
        _fill_cell(row.cells[1], s["workshop_description"])
        _fill_cell(row.cells[2], s["role"])
        _fill_cell(row.cells[3], s["location"])
        _fill_cell(row.cells[4], f"{s['duration_hours']} hrs")
        _fill_cell(row.cells[5], _currency(s["compensation_aed"]))

    # --- Table 1: Add-ons ---
    tbl1 = doc.tables[1]
    orig_tr = tbl1.rows[1]._tr
    row_template = copy.deepcopy(orig_tr)
    tbl1._tbl.remove(orig_tr)
    for a in addons:
        new_tr = copy.deepcopy(row_template)
        tbl1._tbl.append(new_tr)
        row = tbl1.rows[-1]
        _fill_cell(row.cells[0], a["description"])
        _fill_cell(row.cells[1], _currency(a["amount_aed"]))
        _fill_cell(row.cells[2], a.get("notes", ""))

    # Signature block last — it adds a table of its own, so it runs after the
    # doc.tables[0]/[1] lookups above rather than renumbering them.
    _rebuild_signature_block(
        doc,
        [paras[30], paras[31]],
        admin_signatory_name,
        instructor_name,
        signed_date,
        admin_signature_bytes,
        instructor_signature_b64,
    )

    buf = io.BytesIO()
    doc.save(buf)
    return _libreoffice_to_pdf(buf.getvalue())


# ── Excel bulk import (ported from source's parse_excel_bulk_import) ────────

VALID_SESSION_ROLES = {
    "lead facilitator": "Lead Facilitator",
    "facilitator": "Facilitator",
    "assistant facilitator": "Assistant Facilitator",
}


def parse_excel_bulk_import(file_bytes: bytes) -> dict:
    """2-sheet xlsx: 'Sessions' (required) + 'Add-ons' (optional). Returns
    {"instructors": {email: {"sessions": [...], "addons": [...]}}, "errors": [...]}."""
    import openpyxl

    result: dict = {"instructors": {}, "errors": []}
    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    except Exception as e:
        result["errors"].append(f"Cannot open file: {e}")
        return result

    if "Sessions" not in wb.sheetnames:
        result["errors"].append('Missing required sheet named "Sessions".')
        return result

    rows = list(wb["Sessions"].iter_rows(values_only=True))
    if not rows:
        result["errors"].append("Sessions sheet is empty.")
        return result

    headers = [str(h).strip().lower() if h else "" for h in rows[0]]

    def col(name: str) -> int | None:
        try:
            return headers.index(name.lower())
        except ValueError:
            return None

    cols = {
        "email": col("instructor email"), "date": col("session date"), "desc": col("workshop description"),
        "role": col("role"), "location": col("location"), "duration": col("duration (hours)"),
        "comp": col("compensation (aed)"),
    }
    missing = [k for k, v in cols.items() if v is None]
    if missing:
        result["errors"].append(f"Sessions sheet missing columns: {', '.join(missing)}")
        return result

    for row_idx, row in enumerate(rows[1:], start=2):
        email = str(row[cols["email"]]).strip() if row[cols["email"]] else ""
        if not email or email.lower() == "none":
            continue
        try:
            duration = float(row[cols["duration"]]) if row[cols["duration"]] is not None else 0
        except (ValueError, TypeError):
            result["errors"].append(f"Row {row_idx}: Duration must be a number")
            duration = 0
        try:
            compensation = float(row[cols["comp"]]) if row[cols["comp"]] is not None else 0
        except (ValueError, TypeError):
            result["errors"].append(f"Row {row_idx}: Compensation must be a number")
            compensation = 0

        role_raw = str(row[cols["role"]]).strip() if row[cols["role"]] else ""
        role = VALID_SESSION_ROLES.get(role_raw.lower())
        if not role:
            result["errors"].append(f"Row {row_idx}: Invalid role '{role_raw}'")
            role = "Facilitator"

        result["instructors"].setdefault(email, {"sessions": [], "addons": []})["sessions"].append({
            "session_date": str(row[cols["date"]]).strip() if row[cols["date"]] else "",
            "workshop_description": str(row[cols["desc"]]).strip() if row[cols["desc"]] else "",
            "role": role,
            "location": str(row[cols["location"]]).strip() if row[cols["location"]] else "",
            "duration_hours": duration,
            "compensation_aed": compensation,
        })

    if "Add-ons" in wb.sheetnames:
        ao_rows = list(wb["Add-ons"].iter_rows(values_only=True))
        if ao_rows:
            ao_headers = [str(h).strip().lower() if h else "" for h in ao_rows[0]]
            try:
                ao_email, ao_desc, ao_amount = ao_headers.index("instructor email"), ao_headers.index("description"), ao_headers.index("amount (aed)")
                ao_notes = ao_headers.index("notes") if "notes" in ao_headers else None
            except ValueError as e:
                result["errors"].append(f"Add-ons sheet missing column: {e}")
            else:
                for row_idx, row in enumerate(ao_rows[1:], start=2):
                    email = str(row[ao_email]).strip() if row[ao_email] else ""
                    if not email or email.lower() == "none":
                        continue
                    try:
                        amount = float(row[ao_amount]) if row[ao_amount] is not None else 0
                    except (ValueError, TypeError):
                        result["errors"].append(f"Add-ons row {row_idx}: Amount must be a number")
                        amount = 0
                    result["instructors"].setdefault(email, {"sessions": [], "addons": []})["addons"].append({
                        "description": str(row[ao_desc]).strip() if row[ao_desc] else "",
                        "amount_aed": amount,
                        "notes": str(row[ao_notes]).strip() if ao_notes is not None and row[ao_notes] else "",
                    })

    return result


def generate_excel_template() -> bytes:
    import openpyxl
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter  # not cell.column_letter — see instructors/HANDOFF.md known bugs

    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "Sessions"
    headers1 = ["Instructor Email", "Session Date", "Workshop Description", "Role", "Location", "Duration (hours)", "Compensation (AED)"]
    for col_idx, header in enumerate(headers1, 1):
        cell = ws1.cell(row=1, column=col_idx, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="653f84")
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        ws1.column_dimensions[get_column_letter(col_idx)].width = 22
    for col_idx, val in enumerate(
        ["instructor@example.com", "30/10/2025", "Quick Flight Workshop", "Facilitator", "GEMS Metropole School, Dubai", 4, 500], 1
    ):
        ws1.cell(row=2, column=col_idx, value=val)

    ws2 = wb.create_sheet("Add-ons")
    headers2 = ["Instructor Email", "Description", "Amount (AED)", "Notes"]
    for col_idx, header in enumerate(headers2, 1):
        cell = ws2.cell(row=1, column=col_idx, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="653f84")
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        ws2.column_dimensions[get_column_letter(col_idx)].width = 28
    for col_idx, val in enumerate(["instructor@example.com", "Travel allowance", 150, "For session on 30/10"], 1):
        ws2.cell(row=2, column=col_idx, value=val)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
