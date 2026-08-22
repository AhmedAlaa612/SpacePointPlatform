"""Two-column "For SpacePoint FZC / Facilitator" signature block, as a table.

Both signing templates (agreement.docx, payment_letter.docx) lay this block
out inside a single paragraph: the two columns are held apart by literal runs
of padding spaces and counted `<w:tab/>` stops, calibrated against one
particular pair of names. That layout is a function of how wide the names
render, so it only survives names as short as the ones it was measured with.
A Facilitator whose full name ran a few characters long ate the padding, wrapped
the line, and dragged everything after it out of column — the tail of the name
landed under the SpacePoint column and both "Date:" fields ended up floating in
the middle of the page (rendered evidence: a 5-word name moved the SpacePoint
"Date:" from x=72 to x=357 and the Facilitator's from x=396 to x=252).

A table removes the coupling. Each field is its own cell, so a long name wraps
*inside* its own column, the field beside it stays on the same row, and nothing
below it moves. Borders and cell padding are stripped and the columns are equal
halves of the text width, so it prints as the same plain two-column block —
just one that cannot be knocked out of alignment by its contents.

`replace_signature_block` swaps the template's paragraph(s) for that table; the
caller supplies the text and (optional) signature images per column.
"""

import base64
import copy
import io
from dataclasses import dataclass

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt

_ROWS = 4  # heading / Name: / Date: / Signature:
# Only used if a template ever turns up without an anchored signature to
# measure — the size both templates actually ship with, so a signature still
# prints at a sane size instead of raising on a zero width.
_DEFAULT_SIGNATURE_SIZE = (1_514_475, 414_020)


@dataclass(frozen=True)
class SignatureParty:
    """One column of the block. `signature` is raw image bytes (PNG/JPEG)."""

    heading: str
    name: str
    date: str = ""
    signature: bytes | None = None


def replace_signature_block(
    doc,
    paragraphs: list,
    left: SignatureParty,
    right: SignatureParty,
    *,
    signature_size: tuple[int, int],
) -> None:
    """Replace `paragraphs` (the template's signature block) with a 2×4 table.

    The table lands where the first paragraph was; every paragraph in the list
    is then dropped, so pass all of them when a template splits the block over
    more than one (payment_letter.docx keeps the "Signature:" row in its own
    paragraph). `signature_size` is the (cx, cy) EMU size to draw signature
    images at — read it off the template's own anchored signature so the
    images keep the size the document was designed around.
    """
    anchor = paragraphs[0]
    font = _font_rpr(anchor)
    section = doc.sections[0]
    col_w = (section.page_width - section.left_margin - section.right_margin) // 2

    table = doc.add_table(rows=_ROWS, cols=2)
    _strip_table_chrome(table, col_w)

    if not all(signature_size):
        signature_size = _DEFAULT_SIGNATURE_SIZE

    for col, party in ((0, left), (1, right)):
        cells = [table.rows[i].cells[col] for i in range(_ROWS)]
        _fill(cells[0], party.heading, font, col_w, bold=True)
        _fill(cells[1], f"Name: {party.name}".rstrip(), font, col_w)
        _fill(cells[2], f"Date: {party.date}".rstrip(), font, col_w)
        _fill(
            cells[3], "Signature: ", font, col_w,
            picture=party.signature, picture_size=signature_size,
        )
        # The signature image makes its own row taller than a bare "Signature:"
        # label, so bottom-align the row's cells to keep the two labels on the
        # same line when only one party has signed.
        _bottom_align(cells[3])

    # A bare <w:br/> opened the template's block; keep that leading blank line
    # so the table sits where the paragraph did rather than riding up.
    spacer = copy.deepcopy(anchor._p)
    for child in list(spacer):
        if child.tag != qn("w:pPr"):
            spacer.remove(child)

    anchor._p.addnext(table._tbl)
    anchor._p.addnext(spacer)
    for para in paragraphs:
        para._p.getparent().remove(para._p)


def template_signature_image(doc, paragraph) -> tuple[bytes | None, tuple[int, int]]:
    """The signature image anchored on `paragraph`, as (bytes, (cx, cy) EMU).

    Both templates ship the SpacePoint signatory's signature baked in as a
    floating anchor. Pulled out here so the table can redraw it inline in its
    own cell instead of positioning it by hand-measured page offsets.
    """
    run = next(
        (r for r in paragraph._p.findall(qn("w:r")) if r.find(qn("w:drawing")) is not None),
        None,
    )
    if run is None:
        return None, (0, 0)

    extent = run.find(".//" + qn("wp:extent"))
    size = (int(extent.get("cx")), int(extent.get("cy")))

    blip = run.find(".//" + qn("a:blip"))
    rel = blip.get(
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
    )
    part = doc.part.related_parts.get(rel)
    return (part.blob if part is not None else None), size


def decode_signature(signature_b64: str | None) -> bytes | None:
    """Decode a browser-captured signature, with or without a data: prefix."""
    if not signature_b64:
        return None
    raw = signature_b64.split(",", 1)[-1] if "," in signature_b64 else signature_b64
    return base64.b64decode(raw)


def column_headings(text: str, fallback: tuple[str, str]) -> tuple[str, str]:
    """The two column headings, as written in the template.

    Read back rather than hardcoded so the block keeps following the template
    if the wording is ever edited. In both templates the heading line is the
    first one that splits into exactly two pieces across its tab stops.
    """
    for line in text.split("\n"):
        parts = [part.strip() for part in line.split("\t") if part.strip()]
        if len(parts) == 2:
            return parts[0], parts[1]
    return fallback


def _font_rpr(paragraph):
    """The template's own run properties, so the table keeps the document font.

    Taken from the first plain text run — the drawing run carries `<w:noProof/>`
    and the heading runs carry `<w:b/>`, neither of which should apply to every
    cell.
    """
    for run in paragraph.runs:
        if run.text.strip() and not run.bold:
            return run._r.find(qn("w:rPr"))
    return None


def _strip_table_chrome(table, col_w: int) -> None:
    """No borders, no cell padding, fixed equal columns — prints as plain text."""
    tbl_pr = table._tbl.tblPr

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tbl_pr.append(borders)

    margins = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), "0")
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    tbl_pr.append(margins)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    table.autofit = False
    for grid_col in table._tbl.find(qn("w:tblGrid")).findall(qn("w:gridCol")):
        grid_col.set(qn("w:w"), str(Emu(col_w).twips))


def _bottom_align(cell) -> None:
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), "bottom")
    cell._tc.get_or_add_tcPr().append(v_align)


def _fill(cell, text, font_rpr, col_w, *, bold=False, picture=None, picture_size=None):
    cell.width = Emu(col_w)
    para = cell.paragraphs[0]
    # The template block was one paragraph of hard line breaks; the document's
    # 8pt paragraph spacing would otherwise stretch these rows apart.
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)

    run = para.add_run(text)
    if font_rpr is not None:
        existing = run._r.find(qn("w:rPr"))
        if existing is not None:
            run._r.remove(existing)
        run._r.insert(0, copy.deepcopy(font_rpr))
    if bold:
        run.bold = True

    if picture:
        para.add_run().add_picture(
            io.BytesIO(picture), width=Emu(picture_size[0]), height=Emu(picture_size[1])
        )
