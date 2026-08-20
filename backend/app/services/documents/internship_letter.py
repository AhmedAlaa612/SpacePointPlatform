"""Internship letter — docxtpl fills DOCX template → LibreOffice → PDF.
Mirrors services/documents/contract.py's technique exactly (same anchored-
signature-image trick, same LibreOffice conversion), a separate function
since this is a different template/domain (intern, not instructor).

Template: app/static/templates/docx/internship_letter.docx
Placeholders: {{ ref_number }} {{ university_id }} {{ letter_date }}
{{ salutation }} {{ intern_name }} {{ start_date }} {{ duration_weeks }}
{{ activity_description }} {{ hours_per_week }} {{ supervisor_title }}
{{ supervisor_name }} {{ supervisor_first_name }} {{ supervisor_email }}
{{ supervisor_phone }}

The signatory (Abdullah AlSalmani, Co-Founder & CEO) and their signature
image are baked into the template as static content — same convention as
the instructor contract's facilitator identity — since there's exactly one
signatory today and templating it would be speculative.

Signature block (confirmed by inspecting the source docx): a single
paragraph (index 16) reading "Signature:\t...\tIntern's Signature:" with
the admin's signature already anchored in the template; the intern's
signature is anchored the same way on signing, cloning the admin's anchor
XML and shifting only the horizontal offset — same technique as
contract.py's `_add_facilitator_signature_image` / payment_letter.py's
`_inject_signatures`.
"""

import base64
import copy
import io
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docxtpl import DocxTemplate

_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

_TEMPLATE = (
    Path(__file__).parent.parent.parent
    / "static" / "templates" / "docx" / "internship_letter.docx"
)

_SOFFICE = (
    r"C:\Program Files\LibreOffice\program\soffice.exe"
    if sys.platform == "win32"
    else "libreoffice"
)

_SIGNATURE_PARA_IDX = 16
# Measured off the admin's own anchor in the rendered template — the
# intern's signature sits the same horizontal offset past their label as
# the admin's does past "Signature:", mirroring contract.py's
# _FACILITATOR_SIG_OFFSET_H exactly (same tab-heavy-paragraph problem: a
# plain inline picture can wrap to a new line instead of staying in its
# own column, so both signatures are fixed-position anchors).
_INTERN_SIG_OFFSET_H = 4_849_000


def _libreoffice_to_pdf(docx_bytes: bytes) -> bytes:
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "doc.docx"
        src.write_bytes(docx_bytes)
        subprocess.run(
            [_SOFFICE, "--headless", "--convert-to", "pdf", "--outdir", tmp, str(src)],
            check=True, timeout=60, capture_output=True,
        )
        return (Path(tmp) / "doc.pdf").read_bytes()


def _add_intern_signature_image(doc: Document, signature_b64: str) -> None:
    p = doc.paragraphs[_SIGNATURE_PARA_IDX]
    anchor_run = next(
        (r for r in p._p.findall(qn("w:r")) if r.find(qn("w:drawing")) is not None),
        None,
    )
    if anchor_run is None:
        return

    extent = anchor_run.find(".//" + qn("wp:extent"))
    sig_cx, sig_cy = int(extent.get("cx")), int(extent.get("cy"))

    raw = signature_b64.split(",", 1)[-1] if "," in signature_b64 else signature_b64
    from docx.shared import Emu
    tmp = doc.add_paragraph()
    tmp.add_run().add_picture(io.BytesIO(base64.b64decode(raw)), width=Emu(sig_cx), height=Emu(sig_cy))
    blip = tmp._p.find(".//" + qn("a:blip"))
    new_rId = blip.get(f"{{{_R_NS}}}embed")
    tmp._p.getparent().remove(tmp._p)

    intern_run = copy.deepcopy(anchor_run)
    intern_run.find(".//" + qn("a:blip")).set(f"{{{_R_NS}}}embed", new_rId)

    pos_h = intern_run.find(".//" + qn("wp:positionH"))
    pos_h.find(qn("wp:posOffset")).text = str(_INTERN_SIG_OFFSET_H)

    intern_run.find(".//" + qn("wp:docPr")).set("id", "9101")
    intern_run.find(".//" + qn("wp:docPr")).set("name", "intern_sig")
    intern_run.find(".//" + qn("pic:cNvPr")).set("id", "9101")
    intern_run.find(".//" + qn("pic:cNvPr")).set("name", "intern_sig")

    p._p.append(intern_run)


def generate_internship_letter_pdf(
    *,
    ref_number: str,
    university_id: str,
    letter_date: str,
    salutation: str,
    intern_name: str,
    start_date: str,
    duration_weeks: int,
    activity_description: str,
    hours_per_week: int,
    supervisor_title: str,
    supervisor_name: str,
    supervisor_first_name: str,
    supervisor_email: str,
    supervisor_phone: str,
    intern_signature_b64: str | None = None,
) -> bytes:
    tpl = DocxTemplate(str(_TEMPLATE))
    tpl.render({
        "ref_number": ref_number,
        "university_id": university_id,
        "letter_date": letter_date,
        "salutation": salutation,
        "intern_name": intern_name,
        "start_date": start_date,
        "duration_weeks": duration_weeks,
        "activity_description": activity_description,
        "hours_per_week": hours_per_week,
        "supervisor_title": supervisor_title,
        "supervisor_name": supervisor_name,
        "supervisor_first_name": supervisor_first_name,
        "supervisor_email": supervisor_email,
        "supervisor_phone": supervisor_phone,
    })
    buf = io.BytesIO()
    tpl.save(buf)
    buf.seek(0)

    doc = Document(buf)
    if intern_signature_b64:
        _add_intern_signature_image(doc, intern_signature_b64)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return _libreoffice_to_pdf(buf.getvalue())
